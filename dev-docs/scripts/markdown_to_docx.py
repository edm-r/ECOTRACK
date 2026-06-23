#!/usr/bin/env python3
"""Convert a Markdown document to a .docx file with local images.

Supported blocks:
- headings
- paragraphs
- bullet and numbered lists
- tables
- blockquotes
- fenced code blocks
- local images (.png/.jpg/.webp/.gif and .svg)

SVG handling:
- tries CairoSVG if installed
- falls back to macOS Quick Look (`qlmanage`)
"""

from __future__ import annotations

import argparse
import os
import subprocess
import sys
import tempfile
import urllib.parse
import urllib.request
from pathlib import Path
from typing import Iterable

try:
    import markdown
except ImportError as exc:  # pragma: no cover - startup guard
    raise SystemExit(
        "Missing dependency: markdown\n"
        "Install with:\n"
        "  python3 -m pip install markdown"
    ) from exc

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError as exc:  # pragma: no cover - startup guard
    raise SystemExit(
        "Missing dependency: beautifulsoup4\n"
        "Install with:\n"
        "  python3 -m pip install beautifulsoup4"
    ) from exc

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - startup guard
    raise SystemExit(
        "Missing dependency: python-docx\n"
        "Install with:\n"
        "  python3 -m pip install python-docx"
    ) from exc

try:
    from PIL import Image
except ImportError as exc:  # pragma: no cover - startup guard
    raise SystemExit(
        "Missing dependency: Pillow\n"
        "Install with:\n"
        "  python3 -m pip install Pillow"
    ) from exc

try:  # optional
    import cairosvg
except ImportError:  # pragma: no cover - optional
    cairosvg = None


MAX_IMAGE_WIDTH = Inches(6.2)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert a Markdown document to a Word .docx file."
    )
    parser.add_argument(
        "input",
        type=Path,
        help="Path to the source Markdown document.",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="Path to the output .docx file. Defaults to the same name as the input.",
    )
    parser.add_argument(
        "--title",
        help="Optional title stored in Word document properties.",
    )
    return parser.parse_args()


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"


def convert_markdown_to_html(markdown_text: str) -> str:
    return markdown.markdown(
        markdown_text,
        extensions=[
            "extra",
            "tables",
            "fenced_code",
            "sane_lists",
            "nl2br",
        ],
    )


def is_blank_string(node: NavigableString) -> bool:
    return not str(node).strip()


def is_image_only_paragraph(tag: Tag) -> bool:
    meaningful = []
    for child in tag.children:
        if isinstance(child, NavigableString) and is_blank_string(child):
            continue
        meaningful.append(child)
    return len(meaningful) == 1 and isinstance(meaningful[0], Tag) and meaningful[0].name == "img"


def parse_image_source(src: str, markdown_path: Path) -> Path | str:
    parsed = urllib.parse.urlparse(src)
    if parsed.scheme in {"http", "https"}:
        return src
    if parsed.scheme == "file":
        return Path(urllib.request.url2pathname(parsed.path))
    path = Path(src)
    if path.is_absolute():
        return path
    return (markdown_path.parent / path).resolve()


def download_remote_image(url: str) -> Path:
    suffix = Path(urllib.parse.urlparse(url).path).suffix or ".img"
    with urllib.request.urlopen(url) as response:
        data = response.read()
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    Path(tmp_path).write_bytes(data)
    return Path(tmp_path)


def convert_svg_to_png(svg_path: Path) -> Path:
    if cairosvg is not None:
        fd, tmp_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        cairosvg.svg2png(url=str(svg_path), write_to=tmp_path)
        return Path(tmp_path)

    if sys.platform == "darwin" and shutil_which("qlmanage"):
        with tempfile.TemporaryDirectory() as temp_dir:
            subprocess.run(
                [
                    "qlmanage",
                    "-t",
                    "-s",
                    "2200",
                    "-o",
                    temp_dir,
                    str(svg_path),
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            candidates = list(Path(temp_dir).glob(f"{svg_path.name}*.png"))
            if not candidates:
                raise RuntimeError(f"Quick Look did not render {svg_path}")
            fd, tmp_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            Path(tmp_path).write_bytes(candidates[0].read_bytes())
            return Path(tmp_path)

    raise RuntimeError(
        "SVG conversion unavailable. Install cairosvg or run on macOS with qlmanage."
    )


def shutil_which(command: str) -> str | None:
    return subprocess.run(
        ["which", command],
        check=False,
        capture_output=True,
        text=True,
    ).stdout.strip() or None


def image_to_supported_file(src: Path | str) -> tuple[Path, list[Path]]:
    temp_paths: list[Path] = []
    if isinstance(src, str):
        src = download_remote_image(src)
        temp_paths.append(src)

    if src.suffix.lower() == ".svg":
        converted = convert_svg_to_png(src)
        temp_paths.append(converted)
        return converted, temp_paths

    return src, temp_paths


def compute_image_width(image_path: Path):
    try:
        with Image.open(image_path) as image:
            width_px, height_px = image.size
            dpi_x = image.info.get("dpi", (96, 96))[0] or 96
            width_inches = width_px / dpi_x
            return min(Inches(width_inches), MAX_IMAGE_WIDTH)
    except Exception:
        return MAX_IMAGE_WIDTH


def add_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_image(doc: Document, src: str, alt: str, markdown_path: Path) -> None:
    target = parse_image_source(src, markdown_path)

    try:
        prepared, temp_paths = image_to_supported_file(target)

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(prepared), width=compute_image_width(prepared))

        if alt:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption.add_run(alt)
            cap_run.italic = True
            cap_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    except Exception as exc:
        temp_paths = []
        fallback = doc.add_paragraph()
        fallback.style = "Intense Quote" if "Intense Quote" in doc.styles else "Normal"
        fallback.add_run(f"[Image non inseree: {src}] ").bold = True
        fallback.add_run(str(exc))
    finally:
        for path in temp_paths:
            try:
                path.unlink(missing_ok=True)
            except Exception:
                pass


def render_inline_children(paragraph, node: Tag | NavigableString) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    name = node.name or ""
    if name == "br":
        paragraph.add_run().add_break()
        return
    if name == "img":
        alt = node.get("alt", "")
        paragraph.add_run(f"[Image: {alt}]")
        return
    if name == "code" and node.parent and node.parent.name != "pre":
        run = paragraph.add_run(node.get_text())
        run.font.name = "Courier New"
        return
    if name == "a":
        text = node.get_text(" ", strip=False) or node.get("href", "")
        href = node.get("href", "")
        link_run = paragraph.add_run(text)
        link_run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        link_run.underline = True
        if href and href != text:
            suffix = paragraph.add_run(f" ({href})")
            suffix.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        return

    run = None
    if name in {"strong", "b", "em", "i"} and node.get_text():
        run = paragraph.add_run(node.get_text())
        if name in {"strong", "b"}:
            run.bold = True
        if name in {"em", "i"}:
            run.italic = True
        return

    for child in node.children:
        render_inline_children(paragraph, child)


def add_paragraph_from_tag(doc: Document, tag: Tag, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    for child in tag.children:
        render_inline_children(paragraph, child)


def render_list(doc: Document, list_tag: Tag, level: int = 0) -> None:
    style = "List Number" if list_tag.name == "ol" else "List Bullet"
    for li in list_tag.find_all("li", recursive=False):
        paragraph = doc.add_paragraph(style=style)
        if level:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)

        nested_lists: list[Tag] = []
        for child in li.children:
            if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                nested_lists.append(child)
                continue
            render_inline_children(paragraph, child)

        for nested in nested_lists:
            render_list(doc, nested, level + 1)


def render_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return

    matrix: list[list[Tag]] = []
    max_cols = 0
    for row in rows:
        cells = row.find_all(["th", "td"], recursive=False)
        matrix.append(cells)
        max_cols = max(max_cols, len(cells))

    table = doc.add_table(rows=len(matrix), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row in enumerate(matrix):
        for col_idx, cell_tag in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            for child in cell_tag.children:
                render_inline_children(paragraph, child)
            if cell_tag.name == "th":
                for run in paragraph.runs:
                    run.bold = True


def render_code_block(doc: Document, pre_tag: Tag) -> None:
    code_text = pre_tag.get_text("\n")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    add_shading(paragraph, "E5E7EB")
    run = paragraph.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)


def render_blockquote(doc: Document, tag: Tag) -> None:
    paragraph = doc.add_paragraph(style="Intense Quote" if "Intense Quote" in doc.styles else None)
    for child in tag.children:
        render_inline_children(paragraph, child)


def iter_root_nodes(soup: BeautifulSoup) -> Iterable[Tag | NavigableString]:
    root = soup.body if soup.body else soup
    for node in root.children:
        if isinstance(node, NavigableString) and is_blank_string(node):
            continue
        yield node


def render_document(doc: Document, soup: BeautifulSoup, markdown_path: Path) -> None:
    for node in iter_root_nodes(soup):
        if isinstance(node, NavigableString):
            doc.add_paragraph(str(node).strip())
            continue

        if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(node.name[1])
            add_paragraph_from_tag(doc, node, style=f"Heading {min(level, 4)}")
            continue
        if node.name == "p":
            if is_image_only_paragraph(node):
                image = node.find("img")
                add_image(doc, image.get("src", ""), image.get("alt", ""), markdown_path)
            else:
                add_paragraph_from_tag(doc, node)
            continue
        if node.name in {"ul", "ol"}:
            render_list(doc, node)
            continue
        if node.name == "table":
            render_table(doc, node)
            continue
        if node.name == "pre":
            render_code_block(doc, node)
            continue
        if node.name == "blockquote":
            render_blockquote(doc, node)
            continue
        if node.name == "hr":
            doc.add_paragraph("_" * 60)
            continue
        if node.name == "img":
            add_image(doc, node.get("src", ""), node.get("alt", ""), markdown_path)
            continue

        # Fallback for containers such as div/span generated by markdown extensions.
        add_paragraph_from_tag(doc, node)


def convert(markdown_path: Path, output_path: Path, title: str | None) -> None:
    markdown_text = markdown_path.read_text(encoding="utf-8")
    html = convert_markdown_to_html(markdown_text)
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    set_document_defaults(doc)
    doc.core_properties.title = title or markdown_path.stem

    render_document(doc, soup, markdown_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    args = parse_args()
    input_path = args.input.resolve()
    output_path = (
        args.output.resolve()
        if args.output
        else input_path.with_suffix(".docx")
    )

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    convert(input_path, output_path, args.title)
    print(f"Word document generated: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
